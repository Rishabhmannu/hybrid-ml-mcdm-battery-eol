"""
Build TOPIC_PROPOSAL_Short_REVISED.docx from the markdown content.
Run: conda run -n Eco-Research python build_proposal_docx.py
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

title = doc.add_heading('Topic Proposal -- Digital Economics Term Paper', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Course: Digital Economics (Semester 8)   |   Group: [Group Number]   |   Date: April 2026').italic = True

doc.add_heading('1. Topic Title', level=1)
p = doc.add_paragraph()
p.add_run("Digital Product Passports for India's Circular EV Battery Economy: An Integrated ML + MCDM Decision-Support Framework for Dual-Compliance (BWMR + EU Annex XIII) End-of-Life Routing with Indian-Context Synthetic Data").bold = True

t = doc.add_table(rows=4, cols=2)
t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = 'Domain'
t.rows[0].cells[1].text = 'EV & Battery Supply Chains + Digital Product Passport / Traceability + Policy Analysis (India vs. EU)'
t.rows[1].cells[0].text = 'Problem'
t.rows[1].cells[1].text = 'End-of-life routing decisions for retired EV batteries in India are made without data-driven support, creating compliance risk under BWMR (2022/2024/2025) and the EU Digital Product Passport mandate (Feb 2027), while wasting residual battery value. No Indian battery cycling datasets are publicly available.'
t.rows[2].cells[0].text = 'Method'
t.rows[2].cells[1].text = 'Physics-based simulation (PyBaMM) of Indian Driving Cycles for synthetic Indian-context data, combined with hybrid machine learning (XGBoost + LSTM + SHAP) for battery health grading, Fuzzy BWM-TOPSIS for route ranking, and a dual-compliance Digital Passport schema mapping layer.'
t.rows[3].cells[0].text = 'Objective'
t.rows[3].cells[1].text = 'A reproducible, regulator-aligned decision-support framework and reference Digital Passport schema for Indian battery producers and recyclers.'

doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    "India has entered its first wave of EV battery retirements. Most retired batteries still retain 60-80% of their original capacity -- unsuitable for vehicles but perfectly usable for stationary storage (home, grid, telecom). Yet the industry default is immediate recycling, which wastes residual value and fails to meet new regulatory targets."
)
p = doc.add_paragraph()
p.add_run('Two regulations converge in 2027:').bold = True
doc.add_paragraph(
    "India's BWMR (2022/2024/2025) requires battery producers to meet recovery targets scaling to 90% by FY 2026-27, and recycled-content targets rising to 20% by 2030-31.",
    style='List Bullet'
)
doc.add_paragraph(
    "The EU Battery Regulation 2023/1542 mandates a Digital Product Passport (Annex XIII) for every EV battery above 2 kWh placed on the EU market from February 2027 -- affecting every Indian battery exporter.",
    style='List Bullet'
)
doc.add_paragraph(
    "There is currently no framework that combines data-driven battery health prediction with structured end-of-life route selection, while producing compliance-ready digital records for both Indian and EU regulations in a single pipeline. Compounding this gap: no public Indian battery cycling dataset exists, making India-specific research dependent on either proprietary data (unobtainable for academic work in a short timeline) or physics-based simulation of Indian operating conditions."
)

doc.add_heading('3. Research Gap and Contribution', level=1)
doc.add_paragraph(
    "Hybrid machine learning for battery health prediction is well-studied. Multi-criteria decision methods for battery end-of-life are also well-studied. Digital Battery Passport frameworks have been proposed in multiple forms. We acknowledge this crowded methodological field and do not claim novelty in any individual method."
)
p = doc.add_paragraph()
p.add_run('What remains unaddressed ').bold = True
p.add_run(
    "in the literature is the integration of these layers with (i) physics-based simulation of Indian operating conditions to bridge the Indian data gap, (ii) MCDM criteria weights derived from Indian regulatory priorities, and (iii) generation of dual-compliance (BWMR + EU Annex XIII) Digital Passport records in one reproducible framework."
)
p = doc.add_paragraph()
p.add_run('Our contribution: ').bold = True
p.add_run(
    "An integrated, India-specific decision-support framework that (i) bridges the Indian battery data gap using PyBaMM-based physics simulation of ARAI Indian Driving Cycles, (ii) grades batteries using established hybrid ML (XGBoost + LSTM + SHAP), (iii) ranks end-of-life routes using regulatory-priority-weighted Fuzzy BWM-TOPSIS, and (iv) emits a Digital Passport record mapped to both Indian and EU regulatory fields."
)

doc.add_heading('4. Research Questions', level=1)
p = doc.add_paragraph()
p.add_run('RQ1: ').bold = True
p.add_run(
    "How accurately can a hybrid machine learning pipeline (XGBoost for State-of-Health, LSTM for Remaining Useful Life, with SHAP-based feature interpretation) predict battery health across diverse chemistries, operating temperatures, and cycling protocols -- including physics-simulated Indian Driving Cycle conditions?"
)
p = doc.add_paragraph()
p.add_run('RQ2: ').bold = True
p.add_run(
    "How do recommended end-of-life routes change when MCDM weights reflect Indian regulatory priorities (BWMR recovery targets, safety, carbon footprint, EPR economics) versus EU priorities -- and which criteria are the most sensitive drivers of divergence?"
)

doc.add_heading('5. Data', level=1)
doc.add_paragraph(
    "We use four publicly available battery cycling datasets as the primary backbone, supplemented by physics-simulated Indian-context data to address the Indian data gap."
)
doc.add_heading('Real Secondary Data (Public Benchmarks)', level=2)
t = doc.add_table(rows=5, cols=3)
t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = 'Dataset'
t.rows[0].cells[1].text = 'Source & Scale'
t.rows[0].cells[2].text = 'Role'
rows_data = [
    ('BatteryLife', 'HuggingFace (990 batteries, 16 integrated subsets, multiple chemistries; KDD 2025)', 'Primary training & benchmarking'),
    ('Stanford Second-Life', 'Included in BatteryLife ("Stanford_2"; Nissan Leaf retired cells)', 'Second-life validation'),
    ('NASA PCoE', 'NASA Prognostics Center (~34 18650 cells; Randomized set separate)', 'Benchmark comparability'),
    ('CALCE', 'University of Maryland (multiple chemistries)', 'Benchmark comparability'),
]
for i, (a, b, c) in enumerate(rows_data, 1):
    t.rows[i].cells[0].text = a
    t.rows[i].cells[1].text = b
    t.rows[i].cells[2].text = c

doc.add_heading('Physics-Simulated Indian-Context Data (Novel Addition)', level=2)
for line in [
    "Tool: PyBaMM (open-source, peer-reviewed battery simulator; Sulzer et al., J. Open Res. Software, 2021)",
    "Input profiles: ARAI Indian Driving Cycles (IDC, Modified IDC, Indian Urban Driving Cycle) from CMVR/TAP-115/116 (public)",
    "Conditions: Indian ambient 35-45C thermal profiles; 2W/3W high-C-rate duty cycles",
    "Output: ~50-100 virtual cells across NMC and LFP chemistries, with degradation mechanisms (SEI growth, Li plating) enabled",
    "Role: Bridges the Indian battery data gap; enables India-specific validation without proprietary data access",
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('Regulatory Corpus & Policy Data', level=2)
doc.add_paragraph(
    "India BWMR 2022/2024/2025 gazettes, EU Regulation 2023/1542 (Annex XIII fields), Global Battery Alliance Passport data model, CPCB EPR Portal public national dashboard (scraped for India-specific compliance data)."
)
p = doc.add_paragraph()
p.add_run('Limitation: ').bold = True
p.add_run(
    "All four public datasets are US/EU-tested. Indian-context validity is addressed via PyBaMM physics simulation, not primary data collection (infeasible in 30 days without supplier access)."
)
doc.add_heading('MCDM Criteria Weights', level=2)
doc.add_paragraph(
    "Derived from a systematic review of 8-10 peer-reviewed papers on EV battery EoL decision criteria, with BWM consistency ratios reported and sensitivity analysis across 5 weighting scenarios (BWMR-heavy, Annex-XIII-heavy, safety-heavy, carbon-heavy, economics-heavy)."
)

doc.add_heading('6. Methodology', level=1)
doc.add_paragraph(
    "The framework follows an extended pipeline, combining multiple methodological approaches as required by the project guidelines (Machine Learning + Deep Learning + Hybrid MCDM):"
)
p = doc.add_paragraph()
code_run = p.add_run(
    "ARAI Indian Driving Cycles -> PyBaMM simulation --+\n"
    "Public US/EU datasets (BatteryLife + Stanford + NASA + CALCE) --+\n"
    "                                                  v\n"
    "Merged dataset -> Anomaly filter (VAE) -> Feature engineering ->\n"
    "XGBoost (SoH) + LSTM (RUL) -> SHAP explainability ->\n"
    "Grade A/B/C/D -> Fuzzy BWM-TOPSIS route ranking (Indian weights + EU weights) ->\n"
    "Dual-compliance Digital Passport JSON output (BWMR + EU Annex XIII)"
)
code_run.font.name = 'Consolas'
code_run.font.size = Pt(10)

doc.add_heading('Mandatory Baselines for Publishable Rigour', level=2)
doc.add_paragraph("The framework is compared against these baselines. Without them, reviewers will reject for insufficient comparative evaluation.")
baselines = [
    "Simple capacity-fade linear regression (trivial reference)",
    "XGBoost-only State-of-Health prediction (no LSTM, no VAE)",
    "LSTM-only Remaining Useful Life prediction (no XGBoost feature model)",
    "Equal-weight MCDM (no Fuzzy BWM)",
    "Industry's fixed 80% State-of-Health threshold rule (naive routing)",
    "ML trained only on real US data without PyBaMM synthetic Indian augmentation"
]
for b in baselines:
    doc.add_paragraph(b, style='List Number')

doc.add_heading('7. Expected Outcomes', level=1)
outcomes = [
    "A reproducible ML pipeline with publicly released code and benchmark results across chemistries, operating conditions, and Indian-simulated cycles.",
    "Health-grade-based route rankings under both Indian (BWMR-weighted) and EU (Annex-XIII-weighted) criteria, with sensitivity analysis across 5 weighting scenarios.",
    "A reference Digital Passport schema (JSON) mapping framework outputs to BWMR traceability fields and EU Annex XIII fields, with a gap analysis of fields currently unpopulable in Indian practice.",
    "A PyBaMM-based Indian-context synthetic dataset (released openly) as a contribution to the Indian EV research community.",
    "Policy-actionable insights for CPCB on State-of-Health thresholds and recovery targets, and for OEMs on pre-retirement data-collection requirements."
]
for o in outcomes:
    doc.add_paragraph(o, style='List Number')
p = doc.add_paragraph()
p.add_run('Realistic performance targets ').bold = True
p.add_run(
    "(not guarantees): SoH R2 > 0.95 on CALCE/NASA; expected to be lower on BatteryLife diverse-conditions split and PyBaMM Indian-cycle synthetic data. Honest reporting of cross-distribution performance gap."
)

doc.add_heading('8. Fit with Digital Economics', level=1)
doc.add_paragraph("This project sits squarely within digital economics:")
fit_items = [
    ("Digitalization of a physical supply chain", " -- transforming ad-hoc, gut-feel battery routing into a data-driven digital workflow."),
    ("Digital Product Passports as market infrastructure", " -- the DPP is the digital backbone of a new circular-economy market; we generate the data layer that feeds it."),
    ("AI-enabled decision-making", " -- reducing information asymmetry in the second-life battery market, a classic market-design problem."),
    ("Platform potential", " -- the framework is usable by producers, recyclers, and second-life operators, with characteristic digital-platform network effects."),
    ("Data infrastructure for regulatory compliance", " -- directly addresses India's BWMR digital traceability mandate and EU's Digital Battery Passport requirements, both of which are explicitly digital-economics policy instruments."),
]
for title_, rest in fit_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_).bold = True
    p.add_run(rest)

doc.add_heading('9. Publication Plan', level=1)
pubs = [
    ("Primary target: ", "Batteries (MDPI, IF ~4.0) or World Electric Vehicle Journal (MDPI) -- both offer tight methodological and scope fit."),
    ("Fallback: ", "Sustainability (MDPI, IF ~3.3, Q2, ~35-45% acceptance, ~18 days to first decision)."),
    ("Stretch: ", "Resources, Conservation and Recycling (Elsevier, IF ~11.2) -- attempted only if framework performance and novelty exceed review expectations."),
]
for lbl, txt in pubs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(lbl).bold = True
    p.add_run(txt)

doc.add_heading('10. Timeline and Team Allocation', level=1)
doc.add_paragraph("The project is scoped for a 1-month timeline with 10 group members, split into three sub-teams:")
t = doc.add_table(rows=4, cols=3)
t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = 'Sub-team'
t.rows[0].cells[1].text = 'Members'
t.rows[0].cells[2].text = 'Responsibility'
teams = [
    ('ML Pipeline', '4', 'VAE, XGBoost, LSTM, SHAP, mandatory baselines (B1-B6), benchmarking'),
    ('MCDM + Regulatory + Data Engineering', '3', 'Fuzzy BWM-TOPSIS, BWMR/Annex XIII corpus, sensitivity analysis, CPCB EPR portal scraping, PyBaMM Indian-cycle simulation'),
    ('DPP Schema + Writeup', '3', 'Dual-compliance JSON schema design, gap analysis, paper drafting, figures, final coordination'),
]
for i, (a, b, c) in enumerate(teams, 1):
    t.rows[i].cells[0].text = a
    t.rows[i].cells[1].text = b
    t.rows[i].cells[2].text = c
doc.add_paragraph("Weekly WIP will be maintained in the shared Excel sheet as required by the project guidelines.")

doc.add_heading('11. Known Risks and Mitigations', level=1)
t = doc.add_table(rows=5, cols=2)
t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = 'Risk'
t.rows[0].cells[1].text = 'Mitigation'
risks = [
    ("All public datasets are non-Indian", "PyBaMM physics simulation of Indian Driving Cycles addresses this directly"),
    ("Methodological novelty is limited", "Contribution positioned as integration + India-specific encoding + data gap bridge (all genuinely novel combinations)"),
    ("30-day timeline is tight", "Scope frozen early; GAN augmentation held as optional stretch goal; no primary supplier data collection"),
    ("BatteryLife processed version is gated on HuggingFace", "Access application filed on Day 1; CALCE/NASA fallback ready"),
]
for i, (a, b) in enumerate(risks, 1):
    t.rows[i].cells[0].text = a
    t.rows[i].cells[1].text = b

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Submitted for approval by: [Group Members' Names]").bold = True

out_path = '/Users/rishabh/Documents/Eco-Sem8/eco-sem8-group-project/report/TOPIC_PROPOSAL_Short_REVISED.docx'
doc.save(out_path)
print(f"Saved: {out_path}")
